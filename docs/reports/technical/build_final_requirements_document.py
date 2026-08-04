"""Build the privacy-safe final requirements, design, and testing document."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUTPUT = ROOT / "docs" / "reports" / "technical" / "final-requirements-design-testing.docx"
DESKTOP_OUTPUT = Path.home() / "Desktop" / "Final-Requirements-Design-and-Testing.docx"
TEST_PLAN = ROOT / "docs" / "testing" / "testing-plan.docx"
EXTENSION_MD = ROOT / "docs" / "requirements" / "user-story-extensions.md"
CURRENT_TEST_MD = ROOT / "docs" / "testing" / "current-regression-test-cases.md"

NAVY = "17365D"
BLUE = "366091"
MID_BLUE = "4F81BD"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"
WHITE = "FFFFFF"
GRID = "A6B7C8"


BASELINE_STORIES = [
    ("Farmer", "A.1", "As a farmer, I want to upload maize images, so that I can analyse crop conditions."),
    ("Farmer", "A.2", "As a farmer, I want the system to automatically count maize tassels, so that I do not need to perform manual counting."),
    ("Farmer", "A.3", "As a farmer, I want to view counting results clearly, so that I can understand plant growth."),
    ("Farmer", "A.4", "As a farmer, I want to see highlighted tassels on images, so that I can visually verify the results."),
    ("Farmer", "A.5", "As a farmer, I want to upload multiple images at once, so that I can save time."),
    ("Farmer", "A.6", "As a farmer, I want to receive results within a short response time, so that I can make timely decisions."),
    ("Farmer", "A.7", "As a farmer, I want to access the system via mobile devices, so that I can use it in the field."),
    ("Farmer", "A.8", "As a farmer, I want an intuitive and user-friendly interface, so that I can use the system easily."),
    ("Researcher", "B.1", "As a researcher, I want accurate tassel counting results, so that I can conduct reliable analysis."),
    ("Researcher", "B.2", "As a researcher, I want to export results in standard formats, so that I can use them for further research."),
    ("Researcher", "B.3", "As a researcher, I want to analyse historical data, so that I can study trends over time."),
    ("Researcher", "B.4", "As a researcher, I want to compare outputs from different models, so that I can evaluate performance."),
    ("Researcher", "B.5", "As a researcher, I want access to raw datasets, so that I can preprocess and analyse data."),
    ("Researcher", "B.6", "As a researcher, I want to generate visual reports, so that I can present findings effectively."),
    ("Agronomist", "C.1", "As an agronomist, I want to evaluate plant health based on tassel count, so that I can provide recommendations."),
    ("Agronomist", "C.2", "As an agronomist, I want to monitor crop growth over time, so that I can track development stages."),
    ("Agronomist", "C.3", "As an agronomist, I want to detect abnormal patterns in tassel counts, so that I can identify potential issues early."),
    ("Agronomist", "C.4", "As an agronomist, I want a dashboard view of multiple fields, so that I can analyse large-scale crop conditions."),
    ("Agronomist", "C.5", "As an agronomist, I want summarized insights, so that I can make decisions efficiently."),
    ("Admin", "D.1", "As an admin, I want to manage user accounts, so that I can control system access."),
    ("Admin", "D.2", "As an admin, I want to store uploaded images securely, so that data is protected."),
    ("Admin", "D.3", "As an admin, I want to monitor system usage, so that I can ensure system performance."),
    ("Admin", "D.4", "As an admin, I want to manage datasets, so that the system maintains high-quality training data."),
    ("Admin", "D.5", "As an admin, I want to control user permissions, so that different roles have appropriate access levels."),
    ("Admin", "D.6", "As an admin, I want to back up data regularly, so that data loss is prevented."),
    ("System", "E.1", "As a system, I want to preprocess image data, so that the model can achieve better performance."),
    ("System", "E.2", "As a system, I want to train deep learning models, so that tassel counting accuracy can be improved."),
    ("System", "E.3", "As a system, I want to evaluate model performance using appropriate metrics, so that accuracy can be measured and improved."),
    ("System", "E.4", "As a system, I want to deploy the trained model as a service, so that users can access it online."),
    ("System", "E.5", "As a system, I want to support system updates and model improvements, so that new features and enhancements can be integrated."),
]


USE_CASES = [
    ("UC-01", "Upload and count tassels", "Farmer", "Authenticated active account", "Validate and encrypt image; run active tassel model; save count, boxes, confidence, and timing.", "Invalid/oversized image is rejected; unavailable model returns a controlled service error."),
    ("UC-02", "Screen a maize leaf", "Farmer", "Authenticated account and valid JPG/PNG", "Run quality gate and calibrated classifier; format the selected language; persist evidence and response.", "Poor image requests a retake; uncertain evidence is labelled for confirmation."),
    ("UC-03", "Maintain account", "Any signed-in user", "Current password is known", "Update email or password; increment session version; issue a fresh token when appropriate.", "Wrong password, duplicate email, weak password, and stale token are rejected."),
    ("UC-04", "Retrieve governed research data", "Researcher", "Required permission", "Filter history or export an approved dataset with manifest.", "A path outside configured dataset roots is rejected."),
    ("UC-05", "Evaluate a registered model", "Researcher/Admin", "Registered artifact and dataset YAML", "Validate path and digest; run evaluation; persist metrics.", "Missing, altered, pointer-only, or out-of-root artifacts are rejected."),
    ("UC-06", "Review assigned diagnosis", "Agronomist", "Field assignment exists", "Load permitted evidence; confirm, correct, or mark inconclusive; audit the review.", "Unassigned field or invalid review decision is rejected."),
    ("UC-07", "Assign Agronomist", "Admin", "Active Agronomist and existing field", "Create, update, list, or remove Field Assignment; write audit log.", "Inactive/non-Agronomist user is rejected."),
    ("UC-08", "Operate and migrate service", "System/Admin", "Configured secrets, database, and model", "Apply checksum-tracked migrations; start bounded Waitress service; expose safe health state.", "Checksum mismatch, missing secret, database failure, or invalid model prevents unsafe startup."),
]


DATA_DICTIONARY = [
    ("roles", "Role catalogue used by server-side authorization.", "role_id", "users"),
    ("users", "Account, status, permission overrides, and session revocation counter.", "user_id", "roles, images, diagnoses, logs, assignments"),
    ("fields", "Field identity, owner, thresholds, health, and anomaly state.", "field_id", "images, diagnoses, assignments"),
    ("field_assignments", "Explicit many-to-many Agronomist responsibility boundary.", "field_id + agronomist_user_id", "fields, users"),
    ("images", "Validated image metadata with original name, UUID name, digest, dimensions, and owner.", "image_id", "users, fields, image_files, results, diagnoses"),
    ("image_files", "Encrypted original or annotated binary image payload.", "file_id", "images"),
    ("detection_results", "Tassel count, confidence, timing, boxes, and review state.", "result_id", "images"),
    ("disease_diagnoses", "Leaf-screening evidence, uncertainty, context, advice, and expert review.", "diagnosis_id", "users, fields, images"),
    ("datasets", "Governed dataset metadata and approved path.", "dataset_id", "training_runs"),
    ("models", "Versioned model metadata, metrics, path, digest, parent, and lifecycle status.", "model_id", "training_runs"),
    ("training_runs", "Bounded local or external training configuration and outcome.", "run_id", "models, datasets"),
    ("system_logs", "Actor, action, details, and timestamp for auditable operations.", "log_id", "users"),
    ("schema_migrations", "Applied SQL filename, SHA-256 checksum, and timestamp.", "migration_name", "none"),
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), GRID)


def set_column_widths(table, widths) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def format_table(table, widths, header=True, font_size=8.5) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, widths)
    set_table_borders(table)
    if header:
        set_repeat_table_header(table.rows[0])
    for ri, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if ri == 0 and header:
                shade(cell, BLUE)
            elif ri % 2 == 0:
                shade(cell, "F5F8FA")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Cambria"
                    run.font.size = Pt(font_size)
                    if ri == 0 and header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        cell.text = str(text)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = "" if value is None else str(value)
    format_table(table, widths, font_size=font_size)
    return table


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(10)
    for style_name, size, color, before, after in (
        ("Title", 26, NAVY, 0, 15),
        ("Heading 1", 14, BLUE, 24, 6),
        ("Heading 2", 13, MID_BLUE, 14, 4),
        ("Heading 3", 11, MID_BLUE, 10, 3),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Title"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Cambria"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "Counting Maize Tassels in the Wild via Deep Neural Network"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_BLUE)
    footer = section.footer.paragraphs[0]
    add_page_field(footer)
    for run in footer.runs:
        run.font.name = "Cambria"
        run.font.size = Pt(9)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Counting Maize Tassels in the Wild\nvia Deep Neural Network")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Final Requirements, System Design, BCE, UML, ERD, and Testing Documentation")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(MID_BLUE)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Project Reference: FYP-26-S2-7\nDocument Version: 1.0\nDate: 4 August 2026\nLanguage: English").font.size = Pt(11)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.left_indent = Inches(0.55)
    note.paragraph_format.right_indent = Inches(0.55)
    note_run = note.add_run("Privacy note: This copy intentionally contains no student, supervisor, assessor, email, telephone, or other personal information.")
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    entries = [
        "1. Executive Summary and Project Description",
        "2. Requirements Baseline - Original 30 User Stories",
        "3. Extension User Stories and Descriptions",
        "4. Extended System Description and Architecture",
        "5. Use Case Model and Descriptions",
        "6. BCE Model and Descriptions",
        "7. Sequence Models",
        "8. ERD and Data Dictionary",
        "9. Security, Privacy, and Operational Controls",
        "10. Testing Strategy and Historical Baseline",
        "11. Current Regression Test Cases",
        "12. Traceability and Completion Evidence",
        "13. Limitations and Deployment Readiness",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(entry)
    doc.add_paragraph("Page numbers are provided in the footer. Headings are Word-native and appear in the Navigation Pane.")
    doc.add_page_break()


def parse_extensions():
    records = []
    role = None
    current = None
    mode = None
    for raw_line in EXTENSION_MD.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if line.startswith("## ") and line[3:] in {"Farmer", "Researcher", "Agronomist", "Admin", "System"}:
            role = line[3:]
        elif line.startswith("### "):
            if current:
                records.append(current)
            match = re.match(r"###\s+([A-E]\.\d+)\s+(.+)", line)
            current = {"role": role, "id": match.group(1), "title": match.group(2), "acceptance": []} if match else None
            mode = None
        elif current and line.startswith("**User Story:**"):
            current["story"] = line.split("**User Story:**", 1)[1].strip()
            mode = None
        elif current and line.startswith("**Description:**"):
            current["description"] = line.split("**Description:**", 1)[1].strip()
            mode = "description"
        elif current and line.startswith("**Acceptance Criteria:**"):
            inline = line.split("**Acceptance Criteria:**", 1)[1].strip()
            if inline:
                current["acceptance"].append(inline)
            mode = "acceptance"
        elif current and line.startswith("**BCE:**"):
            current["bce"] = line.split("**BCE:**", 1)[1].strip()
            mode = None
        elif current and mode == "acceptance" and re.match(r"\d+\.\s", line):
            current["acceptance"].append(re.sub(r"^\d+\.\s*", "", line))
        elif current and mode == "description" and line:
            current["description"] = (current.get("description", "") + " " + line).strip()
    if current:
        records.append(current)
    return records


def add_figure(doc, path: Path, caption: str, width=6.0) -> None:
    p = doc.add_paragraph(style="Figure Caption")
    p.add_run(caption)
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    shape = picture.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", caption)
    doc.add_paragraph()


def add_historical_test_tables(doc: Document) -> None:
    historical = Document(TEST_PLAN)
    doc.add_heading("10. Testing Strategy and Historical Baseline", level=1)
    doc.add_paragraph(
        "The following material is reproduced from the prior Testing Plan as a historical baseline. Its test-case wording is intentionally unchanged. Statements about Week 10 prototype or mock behavior describe that earlier submission and are superseded by the current regression evidence in Section 11."
    )
    doc.add_heading("10.1 Historical Overview and Environment", level=2)
    for paragraph in historical.paragraphs[3:32]:
        text = paragraph.text.strip()
        if not text:
            continue
        level = 2 if paragraph.style and paragraph.style.name == "Heading 1" else 3 if paragraph.style and paragraph.style.name == "Heading 2" else None
        if level:
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    doc.add_heading("10.2 Historical 17-Case Summary", level=2)
    source = historical.tables[0]
    headers = [cell.text for cell in source.rows[0].cells]
    rows = [[cell.text for cell in row.cells] for row in source.rows[1:]]
    add_table(doc, headers, rows, [0.45, 0.45, 1.25, 1.15, 2.1, 0.6], font_size=7.5)

    doc.add_heading("10.3 Historical Per-Story Test Cases", level=2)
    statuses = [p.text for p in historical.paragraphs if p.text.startswith("Overall Status:")]
    role_counts = [("Farmer", 8), ("Researcher", 6), ("Agronomist", 5), ("Admin", 6), ("System", 5)]
    story_index = 0
    table_index = 1
    for role, count in role_counts:
        doc.add_heading(f"10.3.{role_counts.index((role, count)) + 1} {role}", level=3)
        for _ in range(count):
            meta_source = historical.tables[table_index]
            step_source = historical.tables[table_index + 1]
            test_id = meta_source.rows[0].cells[1].text
            doc.add_heading(test_id, level=3)
            meta_rows = [[cell.text for cell in row.cells] for row in meta_source.rows]
            add_table(doc, ["Field", "Historical text"], meta_rows, [1.35, 4.65], font_size=8)
            step_headers = [cell.text for cell in step_source.rows[0].cells]
            step_rows = [[cell.text for cell in row.cells] for row in step_source.rows[1:]]
            add_table(doc, step_headers, step_rows, [0.4, 1.7, 2.65, 1.25], font_size=7.5)
            if story_index < len(statuses):
                p = doc.add_paragraph()
                run = p.add_run(statuses[story_index])
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(BLUE)
            story_index += 1
            table_index += 2


def parse_current_tests():
    rows = []
    for line in CURRENT_TEST_MD.read_text(encoding="utf-8").splitlines():
        if not line.startswith("| TC-EXT-"):
            continue
        cells = [cell.strip().replace("`", "") for cell in line.strip("|").split("|")]
        if len(cells) == 7:
            rows.append(cells)
    return rows


def build() -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    doc.add_heading("1. Executive Summary and Project Description", level=1)
    doc.add_paragraph(
        "This project delivers a role-based web and mobile system that reads maize images, counts and highlights tassels, and provides cautious bilingual maize leaf-condition screening. The primary objective remains the reduction of labour-intensive manual tassel counting for large-scale agricultural monitoring. The implemented extension adds image-quality feedback, uncertainty-aware guidance, field-scoped Agronomist review, secure account maintenance, encrypted storage, model integrity checks, controlled migrations, and bounded service workloads."
    )
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "The document baselines requirements, explains the current system, records Use Case, BCE, sequence, and ERD designs, and presents historical plus current test evidence. It follows the supplied preliminary technical-report page system while omitting all personal information."
    )
    doc.add_heading("1.2 Human-Centred Product Position", level=2)
    doc.add_paragraph(
        "A photograph has value when the system converts visible evidence into a clear next action. For tassel counting, the image produces a count, confidence, and visual boxes. For leaf screening, the image produces a quality decision, a cautious similarity statement, uncertainty, follow-up questions, and safe next steps. The platform never presents a leaf-screening output as a confirmed diagnosis."
    )
    doc.add_heading("1.3 Roles", level=2)
    add_table(
        doc,
        ["Role", "Primary responsibility"],
        [
            ("Farmer", "Capture/upload field images, receive counts and bilingual guidance, maintain the account."),
            ("Researcher", "Analyse governed results, datasets, provenance, comparisons, and evaluation metrics."),
            ("Agronomist", "Review only assigned fields, interpret evidence, and record accountable decisions."),
            ("Admin", "Manage users, assignments, datasets, models, permissions, storage, backups, and logs."),
            ("System", "Validate content, run separate AI pipelines, encrypt data, bound work, migrate schema, and audit actions."),
        ],
        [1.15, 4.85],
    )

    doc.add_heading("2. Requirements Baseline - Original 30 User Stories", level=1)
    doc.add_paragraph(
        "The 30 User Story sentences in this section are preserved verbatim from the preliminary technical documentation. They are not rewritten, renumbered, or replaced."
    )
    for role in ("Farmer", "Researcher", "Agronomist", "Admin", "System"):
        doc.add_heading(f"2.{('Farmer','Researcher','Agronomist','Admin','System').index(role)+1} {role}", level=2)
        rows = [(story_id, story) for story_role, story_id, story in BASELINE_STORIES if story_role == role]
        add_table(doc, ["ID", "Original User Story"], rows, [0.65, 5.35], font_size=9)

    extensions = parse_extensions()
    doc.add_heading("3. Extension User Stories and Descriptions", level=1)
    doc.add_paragraph(
        "The following 15 stories are additions. Their identifiers continue each role's baseline sequence and make the implemented bilingual, security, review, and operational behavior explicit."
    )
    for role in ("Farmer", "Researcher", "Agronomist", "Admin", "System"):
        role_records = [item for item in extensions if item["role"] == role]
        doc.add_heading(f"3.{('Farmer','Researcher','Agronomist','Admin','System').index(role)+1} {role} Extensions", level=2)
        for item in role_records:
            doc.add_heading(f"{item['id']} {item['title']}", level=3)
            lead = doc.add_paragraph()
            lead.add_run("User Story: ").bold = True
            lead.add_run(item.get("story", ""))
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            p.add_run(item.get("description", ""))
            p = doc.add_paragraph()
            p.add_run("Acceptance Criteria").bold = True
            for criterion in item.get("acceptance", []):
                doc.add_paragraph(criterion.replace("`", ""), style="List Bullet")
            p = doc.add_paragraph()
            p.add_run("BCE Mapping: ").bold = True
            p.add_run(item.get("bce", ""))

    doc.add_heading("4. Extended System Description and Architecture", level=1)
    doc.add_paragraph(
        "The system uses a same-origin responsive PWA and Flask API, PostgreSQL persistence, encrypted image storage, two separate AI inference pipelines, and role-aware review controls. Protected requests reload current account state so that role, status, permission, email, and password changes take effect immediately through session-version revocation."
    )
    doc.add_heading("4.1 Logical Layers", level=2)
    add_table(
        doc,
        ["Layer", "Components", "Responsibility"],
        [
            ("Boundary", "Mobile PWA, desktop pages, Admin console", "Collect images and context; present progress, results, uncertainty, and actions."),
            ("Control", "Auth, upload, inference, advice, assignment, model, migration", "Validate, authorize, orchestrate, bound, and audit every use case."),
            ("Entity", "PostgreSQL and encrypted image payloads", "Persist identities, evidence, results, reviews, models, datasets, and logs."),
            ("Operations", "Waitress, health, backups, migrations", "Start safely, expose readiness, preserve data, and control schema change."),
        ],
        [0.9, 2.0, 3.1],
        font_size=8.5,
    )
    doc.add_heading("4.2 Dual AI Responsibility", level=2)
    doc.add_paragraph(
        "Tassel detection and leaf screening are intentionally separated. The tassel pipeline returns object locations and count evidence. The disease pipeline applies an image-quality gate, calibrated class thresholds, uncertainty states, and a bilingual advice engine. Neither pipeline substitutes mock output when an artifact is unavailable."
    )

    doc.add_heading("5. Use Case Model and Descriptions", level=1)
    add_figure(doc, ROOT / "docs" / "design" / "uml" / "extensions" / "use-case-extended.png", "Figure 1. Extended role-based Use Case model.", width=4.0)
    for use_id, name, actor, precondition, main_flow, alternate in USE_CASES:
        doc.add_heading(f"5.{USE_CASES.index((use_id, name, actor, precondition, main_flow, alternate))+1} {use_id} - {name}", level=2)
        add_table(
            doc,
            ["Item", "Description"],
            [("Primary actor", actor), ("Precondition", precondition), ("Main flow", main_flow), ("Alternate/exception", alternate), ("Postcondition", "Authorized state changes are persisted and an appropriate response is returned.")],
            [1.35, 4.65],
            font_size=8.5,
        )

    doc.add_heading("6. BCE Model and Descriptions", level=1)
    add_figure(doc, ROOT / "docs" / "design" / "uml" / "extensions" / "bce-extended.png", "Figure 2. Consolidated Boundary-Control-Entity model.", width=4.0)
    doc.add_heading("6.1 Boundary Objects", level=2)
    doc.add_paragraph("Mobile PWA and Desktop Web UI serve Farmers, Researchers, and Agronomists. The Admin Console exposes governance functions. Boundaries collect input and present output but do not decide authorization or model trust.")
    doc.add_heading("6.2 Control Objects", level=2)
    add_table(
        doc,
        ["Control", "Key behavior"],
        [
            ("Authentication and Session", "Reload account state; enforce role/permission; reject stale or disabled sessions."),
            ("Upload Validation", "Limit bytes/pixels; verify decode and MIME; generate UUID and digest; encrypt storage."),
            ("Tassel Inference", "Lock model access; use content/model/mode cache key; return boxes and count."),
            ("Disease Assistance", "Gate quality; classify with uncertainty; format bilingual human guidance."),
            ("Field Assignment and Review", "Restrict Agronomist field/evidence access and persist accountable reviews."),
            ("Model and Training Governance", "Allowlist paths, verify digest, bound local work, persist metrics."),
            ("Migration and Startup", "Checksum migrations and refuse unsafe service startup."),
        ],
        [2.0, 4.0],
        font_size=8.5,
    )
    doc.add_heading("6.3 Entity Objects", level=2)
    doc.add_paragraph("Entities retain durable system truth. The ERD in Section 8 defines identifiers and relationships; encrypted Image File data is separated from queryable Image metadata.")

    doc.add_heading("7. Sequence Models", level=1)
    add_figure(doc, ROOT / "docs" / "design" / "uml" / "extensions" / "sequence-tassel-count.png", "Figure 3. Secure mobile upload and tassel-count sequence.", width=6.0)
    doc.add_paragraph("The API authenticates against live database state before image validation. Only validated encrypted evidence receives an image identifier. Inference then uses the active model and persists a Detection Result.")
    add_figure(doc, ROOT / "docs" / "design" / "uml" / "extensions" / "sequence-disease-review.png", "Figure 4. Bilingual disease screening and assigned Agronomist review.", width=6.0)
    doc.add_paragraph("The quality gate can terminate the model flow with retake guidance. When evidence is screenable, model output and user context remain distinct. An Agronomist sees and reviews a diagnosis only through an explicit Field Assignment.")

    doc.add_heading("8. ERD and Data Dictionary", level=1)
    add_figure(doc, ROOT / "docs" / "design" / "database" / "erd-extended.png", "Figure 5. Extended PostgreSQL entity-relationship model.", width=6.0)
    doc.add_heading("8.1 Data Dictionary", level=2)
    add_table(doc, ["Entity", "Purpose", "Primary key", "Principal relationships"], DATA_DICTIONARY, [1.05, 2.45, 1.15, 1.35], font_size=7.5)
    doc.add_heading("8.2 Integrity Rules", level=2)
    for rule in (
        "A User belongs to one Role; account email is unique and session_version revokes stale sessions.",
        "A Field Assignment is unique for one Field/Agronomist pair and is deleted when either parent is removed.",
        "An Image stores metadata and ownership; Image File stores encrypted binary data with one type per image.",
        "Detection Result and Disease Diagnosis are separate outputs because their evidence and review semantics differ.",
        "Model artifact digest and migration checksum make operational changes detectable.",
    ):
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("9. Security, Privacy, and Operational Controls", level=1)
    add_table(
        doc,
        ["Risk", "Implemented control", "Verification"],
        [
            ("Filename collision / disguised image", "UUID name, SHA-256, MIME/decode/pixel validation, request limits", "Image-security and PostgreSQL upload tests"),
            ("Stale privilege token", "Live user reload and session_version; no query token by default", "Farmer-account authentication tests"),
            ("Unrelated field evidence", "Explicit Field Assignment on fields, diagnoses, reviews, and image access", "Field-authorization tests"),
            ("Untrusted model/dataset path", "Configured roots, suffix/type checks, Git LFS pointer rejection, digest verification", "Security-control tests"),
            ("Resource exhaustion", "Rate limits, bounded LRU cache, locked inference, bounded training queue, Waitress threads", "Rate/cache/executor tests"),
            ("Schema drift", "Numbered checksum-tracked migrations under PostgreSQL advisory lock", "Migration discovery and live check"),
            ("Sensitive data disclosure", "Fernet encrypted bytes, controlled file endpoints, generic public errors, audit logs", "Secure round-trip and response tests"),
        ],
        [1.4, 3.2, 1.4],
        font_size=7.5,
    )
    doc.add_heading("9.1 Privacy Boundary", level=2)
    doc.add_paragraph("Farmers access their own evidence. Agronomists access only explicitly assigned fields. Researchers receive role-approved research records and datasets. Admin access is broad but authenticated, permission-controlled, and audited. Tokens are transported in the Authorization header and are not embedded in protected image URLs.")
    doc.add_heading("9.2 Scale Boundary", level=2)
    doc.add_paragraph("The in-memory limiter and cache are appropriate to a single-process academic deployment. A horizontally scaled public deployment must use a shared external rate limiter/cache, managed object storage, centralized monitoring, secret rotation, and automated backup verification.")

    add_historical_test_tables(doc)

    doc.add_heading("11. Current Regression Test Cases", level=1)
    doc.add_paragraph("The current suite was executed after the implementation changes. Result: 78 tests passed in 11.695 seconds on 4 August 2026. The handled persistence-failure test intentionally writes one error log while returning the completed assessment with a failed-persistence status.")
    current_rows = parse_current_tests()
    add_table(doc, ["ID", "Story", "Level", "Preconditions", "Test action", "Expected result", "Evidence"], current_rows, [0.9, 0.45, 0.65, 1.0, 1.05, 1.25, 0.7], font_size=6.8)
    doc.add_heading("11.1 Test Levels", level=2)
    add_table(
        doc,
        ["Level", "Purpose", "Representative coverage"],
        [
            ("Unit", "Validate deterministic policy and transformation logic.", "Image validation, advice isolation, path controls, caches, rate limits."),
            ("API contract", "Verify role, request, status-code, and response-shape behavior.", "Authentication, disease screening/review, model and report controls."),
            ("Database integration", "Verify constraints, encryption metadata, session revocation, and cleanup.", "PostgreSQL upload and Farmer account suites."),
            ("AI artifact", "Verify a real loadable artifact follows the backend contract.", "Temporary TorchScript disease model and configured tassel predictor."),
            ("PWA/static", "Prevent mobile regression and insecure token transport.", "Responsive assets, protected fetch, bilingual interface markers."),
            ("Migration/operation", "Verify reproducible schema and bounded execution.", "Migrations 001-004, training executor, Waitress configuration."),
        ],
        [1.0, 2.2, 2.8],
        font_size=8,
    )

    doc.add_heading("12. Traceability and Completion Evidence", level=1)
    trace_rows = []
    for item in extensions:
        story = item["id"]
        evidence = {
            "A.9": "Disease API, advice engine, disease tests",
            "A.10": "Profile/password APIs, session tests",
            "A.11": "PWA upload client, image isolation tests",
            "B.7": "Result/model metadata and contract tests",
            "B.8": "Dataset path control and export API",
            "B.9": "Model path/digest validation and evaluation API",
            "C.6": "Field Assignment migration, scoped routes, auth tests",
            "C.7": "Diagnosis review API and review contract tests",
            "C.8": "Advice context isolation and safety wording tests",
            "D.7": "Field assignment Admin API and audit log",
            "D.8": "Live session validation and stale-token tests",
            "D.9": "Storage/model/migration/workload security suite",
            "E.6": "Image-security module and validation tests",
            "E.7": "Separate tassel/disease predictors and AI tests",
            "E.8": "Bounded cache/executor, migration runner, Waitress",
        }[story]
        trace_rows.append((story, item["role"], item["title"], evidence, "Implemented and tested"))
    add_table(doc, ["Story", "Role", "Capability", "Code/test evidence", "Status"], trace_rows, [0.55, 0.8, 1.65, 2.25, 0.75], font_size=7.2)
    doc.add_heading("12.1 Git Evidence", level=2)
    doc.add_paragraph("Core implementation commits: 98d34af (secure data/model flows) and 9b38155 (safe startup controls). Diagram PNGs and editable Mermaid sources are versioned together. The original 30 BCE and sequence sources remain in place.")

    doc.add_heading("13. Limitations and Deployment Readiness", level=1)
    doc.add_heading("13.1 Scientific Limitations", level=2)
    doc.add_paragraph("Leaf screening is limited to configured classes and available training data. It is not a laboratory diagnosis and must be validated across geography, crop stage, imaging conditions, devices, and disease prevalence. Treatment decisions require qualified Agronomist review and locally approved guidance.")
    doc.add_heading("13.2 Public Deployment Requirements", level=2)
    for item in (
        "Managed HTTPS and domain with secure reverse-proxy headers.",
        "Strong SECRET_KEY and FILE_ENCRYPTION_KEY stored in a managed secret service.",
        "Managed PostgreSQL with automated encrypted backups and restoration drills.",
        "Durable object storage for image payloads and a shared rate limiter/cache for multiple instances.",
        "Monitoring for latency, queue capacity, storage growth, model availability, and failed reviews.",
        "Versioned model approval with artifact digest, validation report, rollback plan, and Agronomist sign-off.",
    ):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("13.3 Conclusion", level=2)
    doc.add_paragraph("The codebase now supports both core maize tassel counting and human-centred maize leaf screening through explicit role and data boundaries. The original requirements remain traceable, while the extension stories document implemented behavior that is secure, bilingual, testable, and aligned with practical field use.")

    core = doc.core_properties
    core.title = "Final Requirements, Design, and Testing Documentation"
    core.subject = "Counting Maize Tassels in the Wild via Deep Neural Network"
    core.author = ""
    core.last_modified_by = ""
    core.comments = "Privacy-safe academic project documentation"
    update_fields = doc.settings.element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        doc.settings.element.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.save(OUTPUT)
    shutil.copy2(OUTPUT, DESKTOP_OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
